from pathlib import Path
from typing import Dict
from datetime import datetime
from PyQt6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTextEdit, QTableWidget, QTableWidgetItem,
    QFileDialog, QProgressBar, QSpinBox, QComboBox, QGroupBox, QGridLayout, 
    QSplitter, QMessageBox, QHeaderView, QMenu, QToolButton)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QAction
from PyQt6.QtGui import QFont, QColor
from Processing_Thread import ProcessingThread
from Candidate_profile import CandidateProfile
from docx import Document


class ResumeClassifierGUI(QMainWindow):
    
    def __init__(self):
        super().__init__()
        self.candidates = []
        self.filtered_candidates = []
        self.processing_thread = None
        self.init_ui()
    
    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle("Intelligent Resume Classifier")
        self.setGeometry(100, 100, 1400, 900)
        
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        
        # Create splitter for resizable panels
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)
        
        # Left panel - Filters and controls
        left_panel = self.create_left_panel()
        splitter.addWidget(left_panel)
        
        # Right panel - Results
        right_panel = self.create_right_panel()
        splitter.addWidget(right_panel)
        
        # Set splitter proportions
        splitter.setSizes([400, 1000])
        
        # Status bar
        self.statusBar().showMessage("Ready to process resumes")
        
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin-top: 1ex;
                padding-top: 15px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
            QToolButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QToolButton:hover {
                background-color: #45a049;
            }
            QToolButton:pressed {
                background-color: #3d8b40;
            }
            QToolButton::menu-indicator {
                image: none;
                width: 0px;
            }
        """)
    
    def create_left_panel(self) -> QWidget:
        """Create the left panel with filters and controls"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # File selection group
        file_group = QGroupBox("Resume Files")
        file_layout = QVBoxLayout(file_group)
        
        self.select_files_btn = QPushButton("Select Resume Files")
        self.select_files_btn.clicked.connect(self.select_files)
        file_layout.addWidget(self.select_files_btn)
        
        self.selected_files_label = QLabel("No files selected")
        self.selected_files_label.setWordWrap(True)
        file_layout.addWidget(self.selected_files_label)
        
        layout.addWidget(file_group)
        
        # Filtering criteria group
        criteria_group = QGroupBox("Filtering Criteria")
        criteria_layout = QGridLayout(criteria_group)
        
        # Age filter
        criteria_layout.addWidget(QLabel("Age Range:"), 0, 0)
        self.min_age_spin = QSpinBox()
        self.min_age_spin.setRange(18, 70)
        self.min_age_spin.setValue(18)
        criteria_layout.addWidget(QLabel("Min:"), 0, 1)
        criteria_layout.addWidget(self.min_age_spin, 0, 2)
        
        self.max_age_spin = QSpinBox()
        self.max_age_spin.setRange(18, 70)
        self.max_age_spin.setValue(65)
        criteria_layout.addWidget(QLabel("Max:"), 0, 3)
        criteria_layout.addWidget(self.max_age_spin, 0, 4)
        
        # Location filter
        criteria_layout.addWidget(QLabel("Location:"), 1, 0)
        self.location_input = QLineEdit()
        self.location_input.setPlaceholderText("Enter preferred location")
        criteria_layout.addWidget(self.location_input, 1, 1, 1, 4)
        
        # Role filter
        criteria_layout.addWidget(QLabel("Role:"), 2, 0)
        self.role_input = QLineEdit()
        self.role_input.setPlaceholderText("Enter desired role")
        criteria_layout.addWidget(self.role_input, 2, 1, 1, 4)
        
        # Education filter
        criteria_layout.addWidget(QLabel("Education:"), 3, 0)
        self.education_input = QLineEdit()
        self.education_input.setPlaceholderText("Enter education keyword (e.g., Engineering, Medicine)")
        criteria_layout.addWidget(self.education_input, 3, 1, 1, 4)
        
        # Nationality filter
        criteria_layout.addWidget(QLabel("Nationality:"), 4, 0)
        self.nationality_input = QLineEdit()
        self.nationality_input.setPlaceholderText("Enter preferred nationality")
        criteria_layout.addWidget(self.nationality_input, 4, 1, 1, 4)
        
        layout.addWidget(criteria_group)
        
        # Processing controls
        controls_group = QGroupBox("Processing")
        controls_layout = QVBoxLayout(controls_group)
        
        self.process_btn = QPushButton("Process Resumes")
        self.process_btn.clicked.connect(self.process_resumes)
        self.process_btn.setEnabled(False)
        controls_layout.addWidget(self.process_btn)
        
        self.progress_bar = QProgressBar()
        controls_layout.addWidget(self.progress_bar)
        
        self.clear_btn = QPushButton("Clear Results")
        self.clear_btn.clicked.connect(self.clear_results)
        controls_layout.addWidget(self.clear_btn)
        
        layout.addWidget(controls_group)
        
        # Export controls
        export_group = QGroupBox("Export")
        export_layout = QVBoxLayout(export_group)
        
        # Create export dropdown button
        self.export_all_btn = QToolButton()
        self.export_all_btn.setText("Export All CVs")
        self.export_all_btn.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        
        # Create menu for export options
        export_menu = QMenu(self.export_all_btn)
        
        # Add export actions
        export_docx_action = QAction("Export as DOCX", self)
        export_docx_action.triggered.connect(lambda: self.export_to_format('docx'))
        export_menu.addAction(export_docx_action)
        
        export_pdf_action = QAction("Export as PDF", self)
        export_pdf_action.triggered.connect(lambda: self.export_to_format('pdf'))
        export_menu.addAction(export_pdf_action)
        
        self.export_all_btn.setMenu(export_menu)
        self.export_all_btn.setDefaultAction(export_docx_action)  # Default to DOCX
        export_layout.addWidget(self.export_all_btn)

        
        
        layout.addWidget(export_group)
        
        layout.addStretch()
        return panel
    
    def create_right_panel(self) -> QWidget:
        """Create the right panel with results"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # Results header
        header_layout = QHBoxLayout()
        self.results_label = QLabel("Results: 0 candidates processed")
        self.results_label.setFont(QFont("Arial", 12, QFont.Weight.Bold))
        header_layout.addWidget(self.results_label)
        
        header_layout.addStretch()
        
        # Sort options
        sort_label = QLabel("Sort by:")
        header_layout.addWidget(sort_label)
        
        self.sort_combo = QComboBox()
        self.sort_combo.addItems(["Match Score", "Name", "Age", "Nationality"])
        self.sort_combo.currentTextChanged.connect(self.sort_results)
        header_layout.addWidget(self.sort_combo)
        
        layout.addLayout(header_layout)
        
        # Results table
        self.results_table = QTableWidget()
        self.results_table.setColumnCount(7)  # Match Score, Name, Age, Education, Current Role, Location, Nationality
        self.results_table.setHorizontalHeaderLabels(["Match Score", "Name", "Age", "Education", "Current Role", "Location", "Nationality"])
    
        # Set column widths
        header = self.results_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        header.resizeSection(0, 100)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        header.resizeSection(2, 60)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)  # Education column
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Stretch)  # Current Role column
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.Stretch)  # Location column
        header.setSectionResizeMode(6, QHeaderView.ResizeMode.Fixed)    # Nationality column
        header.resizeSection(6, 100)

        self.results_table.setAlternatingRowColors(True)
        self.results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.results_table.itemSelectionChanged.connect(self.show_candidate_details)
        
        layout.addWidget(self.results_table)
        
        # Candidate details
        details_group = QGroupBox("Candidate Details")
        details_layout = QVBoxLayout(details_group)
        
        # Add horizontal layout for buttons
        button_layout = QHBoxLayout()
        
        # Add existing details text
        self.details_text = QTextEdit()
        self.details_text.setReadOnly(True)
        
        # Add to layout
        details_layout.addLayout(button_layout)
        details_layout.addWidget(self.details_text)
        
        layout.addWidget(details_group)
        
        return panel
    
    def select_files(self):
        """Open file dialog to select resume files"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Select Resume Files",
            "",
            "Resume Files (*.pdf *.docx *.txt);;PDF Files (*.pdf);;Word Files (*.docx);;Text Files (*.txt)"
        )
        
        if file_paths:
            self.selected_files = file_paths
            file_count = len(file_paths)
            self.selected_files_label.setText(f"{file_count} files selected")
            self.process_btn.setEnabled(True)
        else:
            self.selected_files = []
            self.selected_files_label.setText("No files selected")
            self.process_btn.setEnabled(False)
    
    def get_filter_criteria(self) -> Dict:
        """Get current filter criteria from UI"""
        criteria = {
            'min_age': self.min_age_spin.value(),
            'max_age': self.max_age_spin.value(),
            'location': self.location_input.text().strip(),
            'role': self.role_input.text().strip(),
            'education': self.education_input.text().strip(),
            'nationality': self.nationality_input.text().strip(),
        }
        
        return criteria
    
    def process_resumes(self):
        """Start processing selected resumes"""
        if not hasattr(self, 'selected_files') or not self.selected_files:
            QMessageBox.warning(self, "Warning", "Please select resume files first.")
            return
        
        criteria = self.get_filter_criteria()
        
        # Disable process button and clear previous results
        self.process_btn.setEnabled(False)
        self.clear_results()
        self.progress_bar.setValue(0)
        
        # Start processing thread
        self.processing_thread = ProcessingThread(self.selected_files, criteria)
        self.processing_thread.progress_updated.connect(self.progress_bar.setValue)
        self.processing_thread.resume_processed.connect(self.add_candidate)
        self.processing_thread.processing_finished.connect(self.processing_complete)
        self.processing_thread.start()
        
        self.statusBar().showMessage("Processing resumes...")
    
    def add_candidate(self, profile: CandidateProfile):
        """Add processed candidate to results"""
        self.candidates.append(profile)
        self.update_results_display()
    
    def processing_complete(self):
        """Handle processing completion"""
        self.process_btn.setEnabled(True)
        self.statusBar().showMessage(f"Processing complete. {len(self.candidates)} candidates processed.")
        self.sort_results()
    
    def update_results_display(self):
        """Update the results table display"""
        self.results_label.setText(f"Results: {len(self.candidates)} candidates processed")
        
        # Update table
        self.results_table.setRowCount(len(self.candidates))
        
        for row, candidate in enumerate(self.candidates):
            # Match Score
            score_item = QTableWidgetItem(f"{candidate.match_score:.2f}")
            score_item.setData(Qt.ItemDataRole.UserRole, candidate.match_score)
            self.results_table.setItem(row, 0, score_item)
            
            # Color code based on match score
            if candidate.match_score >= 0.8:
                score_item.setBackground(QColor(144, 238, 144))  # Light green
            elif candidate.match_score >= 0.6:
                score_item.setBackground(QColor(255, 255, 144))  # Light yellow
            elif candidate.match_score >= 0.4:
                score_item.setBackground(QColor(255, 200, 144))  # Light orange
            else:
                score_item.setBackground(QColor(255, 144, 144))  # Light red
            
            # Other fields
            self.results_table.setItem(row, 1, QTableWidgetItem(candidate.name))
            self.results_table.setItem(row, 2, QTableWidgetItem(str(candidate.age) if candidate.age else "N/A"))
            self.results_table.setItem(row, 3, QTableWidgetItem(candidate.education))
            self.results_table.setItem(row, 4, QTableWidgetItem(candidate.current_role))
            self.results_table.setItem(row, 5, QTableWidgetItem(candidate.current_residence))
            self.results_table.setItem(row, 6, QTableWidgetItem(candidate.nationality))
    
    def sort_results(self):
        """Sort results based on selected criteria"""
        sort_by = self.sort_combo.currentText()
        
        if sort_by == "Match Score":
            self.candidates.sort(key=lambda x: x.match_score, reverse=True)
        elif sort_by == "Name":
            self.candidates.sort(key=lambda x: x.name.lower())
        elif sort_by == "Age":
            self.candidates.sort(key=lambda x: x.age if x.age else 0)
        elif sort_by == "Nationality":
            self.candidates.sort(key=lambda x: x.nationality.lower() if x.nationality else "")
        
        self.update_results_display()
    
    def show_candidate_details(self):
        """Show detailed information for selected candidate"""
        current_row = self.results_table.currentRow()
        if current_row >= 0 and current_row < len(self.candidates):
            candidate = self.candidates[current_row]
            details = f"""
                <b>Name:</b> {candidate.name}<br>
                <b>Age:</b> {candidate.age}<br>
                <b>Education:</b> {candidate.education}<br>
                <b>Current Role:</b> {candidate.current_role}<br>
                <b>Location:</b> {candidate.current_residence}<br>
                <b>Nationality:</b> {candidate.nationality}<br>
                <b>Email:</b> {candidate.email}<br>
                <b>Phone:</b> {candidate.phone}<br>
                <b>Match Score:</b> {candidate.match_score:.2f}<br>
                <b>File:</b> {Path(candidate.file_path).name}<br><br>
                <b>Resume Preview:</b><br>
                {candidate.raw_text}
                """
            self.details_text.setHtml(details)
    
    def clear_results(self):
        """Clear all results"""
        self.candidates.clear()
        self.results_table.setRowCount(0)
        self.details_text.clear()
        self.results_label.setText("Results: 0 candidates processed")
        self.progress_bar.setValue(0)
                
    def export_to_format(self, format_type):
        """Export filtered candidates to specified format"""
        if not self.candidates:
            QMessageBox.information(self, "Info", "No candidates to export.")
            return
        
        # Select directory to save files
        format_upper = format_type.upper()
        output_dir = QFileDialog.getExistingDirectory(
            self,
            f"Select Output Directory for CV Files ({format_upper})",
            ""
        )
        
        if not output_dir:
            return
        
        try:
            if format_type == 'docx':
                self.export_to_docx(output_dir)
            elif format_type == 'pdf':
                self.export_to_pdf(output_dir)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export CVs as {format_upper}: {str(e)}")

    def export_to_docx(self, output_dir):
        """Export filtered candidates to individual DOCX files using the template"""
        template_path = "temp.docx"
        if not Path(template_path).exists():
            QMessageBox.critical(self, "Error", f"Template file not found: {template_path}")
            return
        
        exported_count = 0
        
        for candidate in self.candidates:
            if candidate.match_score >= 0.4:  # Only export candidates with decent match scores
                try:
                    # Create a copy of the template
                    doc = Document(template_path)
                    
                    # Fill the template with candidate data
                    self.fill_template_with_candidate_data(doc, candidate)
                    
                    # Save the filled template
                    safe_name = "".join(c for c in candidate.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    if not safe_name:
                        safe_name = f"Candidate_{exported_count + 1}"
                    
                    output_file = Path(output_dir) / f"{safe_name}_CV.docx"
                    doc.save(str(output_file))
                    exported_count += 1
                    
                except Exception as e:
                    print(f"Error exporting {candidate.name}: {e}")
                    continue
        
        QMessageBox.information(self, "Success", 
                            f"Successfully exported {exported_count} CV files to {output_dir}")

    def export_to_pdf(self, output_dir):
        """Export filtered candidates to individual PDF files"""
        try:
            from docx2pdf import convert
            import tempfile
            import os
        except ImportError:
            QMessageBox.critical(self, "Error", 
                               "PDF export requires 'docx2pdf' package. Please install it using: pip install docx2pdf")
            return
        
        template_path = "temp.docx"
        if not Path(template_path).exists():
            QMessageBox.critical(self, "Error", f"Template file not found: {template_path}")
            return
        
        exported_count = 0
        
        # Create temporary directory for DOCX files
        with tempfile.TemporaryDirectory() as temp_dir:
            for candidate in self.candidates:
                if candidate.match_score >= 0.4:  # Only export candidates with decent match scores
                    try:
                        # Create a copy of the template
                        doc = Document(template_path)
                        
                        # Fill the template with candidate data
                        self.fill_template_with_candidate_data(doc, candidate)
                        
                        # Save the filled template to temp directory
                        safe_name = "".join(c for c in candidate.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                        if not safe_name:
                            safe_name = f"Candidate_{exported_count + 1}"
                        
                        temp_docx = Path(temp_dir) / f"{safe_name}_CV.docx"
                        doc.save(str(temp_docx))
                        
                        # Convert to PDF
                        output_pdf = Path(output_dir) / f"{safe_name}_CV.pdf"
                        convert(str(temp_docx), str(output_pdf))
                        
                        exported_count += 1
                        
                    except Exception as e:
                        print(f"Error exporting {candidate.name} to PDF: {e}")
                        continue
        
        QMessageBox.information(self, "Success", 
                            f"Successfully exported {exported_count} PDF files to {output_dir}")

    def fill_template_with_candidate_data(self, doc, candidate):
        """Fill the template document with candidate data"""
        # Get all tables in the document
        if not doc.tables:
            return
            
        table = doc.tables[0]  # Assuming the main table is the first one
        
        # Helper function to set cell text safely
        def set_cell_text(row_idx, col_idx, text):
            try:
                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    # Clear existing content
                    for paragraph in cell.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)
                    # Add new text
                    cell.add_paragraph(str(text) if text else "")
            except Exception as e:
                print(f"Error setting cell [{row_idx}][{col_idx}]: {e}")
        
        # Fill the template based on the table structure
        # Row 1: Proposed Position
        if candidate.current_role:
            set_cell_text(1, 4, candidate.current_role)
        
        # Row 2: Location
        if candidate.current_residence:
            set_cell_text(2, 4, candidate.current_residence)
        
        # Row 3: Name
        if candidate.name:
            set_cell_text(3, 4, candidate.name)
        
        # Row 4: Date of Birth (calculate from age)
        if candidate.age and candidate.age > 0:
            birth_year = datetime.now().year - candidate.age
            set_cell_text(4, 4, f"Approximately {birth_year}")
        
        # Row 5: Nationality
        if candidate.nationality:
            set_cell_text(5, 4, candidate.nationality)
        
        # Row 7: Education
        if candidate.education:
            set_cell_text(7, 4, candidate.education)
        
        # Row 10: Countries of Work Experience
        if candidate.current_residence:
            set_cell_text(10, 4, candidate.current_residence)
        
        # Row 12: Employment Record - Current position
        if candidate.current_role:
            # Find the employment record rows and fill current position
            for i in range(12, min(20, len(table.rows))):
                try:
                    if table.rows[i].cells[1].text.strip() == "Present":
                        set_cell_text(i, 6, candidate.current_role)  # Role/Position Held
                except Exception as e:
                    continue